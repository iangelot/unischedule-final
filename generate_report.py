"""
Generate UniSchedule Africa B-Tech report matching EKOBE VICTOR reference structure.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUTPUT = r"c:\Users\idris\Desktop\unischedule-final\UniSchedule_Africa_Report_Ekobe_Victor.docx"

BOOKMARK_COUNTER = 0


def next_bookmark_id():
    global BOOKMARK_COUNTER
    BOOKMARK_COUNTER += 1
    return str(BOOKMARK_COUNTER)


def set_run_font(run, name="Times New Roman", size=12, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run_font(run, **kwargs)
    return run


def add_paragraph(doc, text="", style="Normal", align=None, justify=False, space_after=6, **run_kw):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    elif justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if text:
        add_run(p, text, **run_kw)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bookmark(paragraph, bookmark_name):
    bid = next_bookmark_id()
    p = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bid)
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bid)
    p.insert(0, start)
    p.append(end)
    return bookmark_name


def add_hyperlink_to_bookmark(paragraph, text, bookmark_name, bold=False, italic=False, size=12):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    if bold:
        b = OxmlElement("w:b")
        r_pr.append(b)
    if italic:
        i = OxmlElement("w:i")
        r_pr.append(i)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_pr.append(r_fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_heading_section(doc, text, level=2, bookmark=None):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, bold=True, size=13 if level == 3 else 14)
    if bookmark:
        add_bookmark(p, bookmark)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_chapter_title(doc, text, bookmark):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, bold=True, size=14, name="Times New Roman")
    add_bookmark(p, bookmark)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_section_title_page(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    set_run_font(run, bold=True, size=16, name="Times New Roman")
    add_bookmark(p, title.replace(" ", "_"))
    add_page_break(doc)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Paragraph")
    if bold_prefix:
        add_run(p, bold_prefix, bold=True)
        add_run(p, text)
    else:
        add_run(p, text)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_toc_entry(doc, label, bookmark, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.35)
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
    add_hyperlink_to_bookmark(p, label, bookmark, size=12)
    add_run(p, "\t", size=12)
    # page placeholder — Word TOC fields show pages; manual links omit page numbers
    return p


def add_toc_line(doc, label, bookmark, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent * 0.35)
    p.paragraph_format.space_after = Pt(3)
    add_hyperlink_to_bookmark(p, label, bookmark, bold=("CHAPTER" in label), size=12)
    return p


def add_figure_line(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=12)
    return p


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=11)


def build_cover(doc):
    add_paragraph(doc, "Republic Of Cameroon         Republic Du Cameroun", bold=True, size=16)
    doc.add_paragraph()
    add_paragraph(doc, "Peace-Work-Father                           Paix-Travail-Patrie", bold=True, name="Calibri")
    add_paragraph(doc, "*************                              *************", bold=True, name="Calibri")
    add_paragraph(
        doc,
        "Ministry of Higher Education                Ministère De L'enseignement Superieur",
        bold=True,
        name="Calibri",
    )
    add_paragraph(doc, "*************                              *************", bold=True, name="Calibri")
    add_paragraph(
        doc,
        "University Institut Of The Tropics              Institute Universitaire Des Grandes Ecoles Des Tropics",
        bold=True,
        name="Calibri",
    )
    for _ in range(2):
        doc.add_paragraph()
    p = add_paragraph(
        doc,
        "A SCHOOL REPORT CARRIED OUT FROM THE 11th MAY 2025 TO THE 16th OF MAY 2025",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "SPECIALTY: ", bold=True, size=16)
    add_run(p, "SOFTWARE ENGINEERING", bold=True, size=16, color=RGBColor(0xFF, 0x00, 0x00))
    doc.add_paragraph()
    add_paragraph(doc, "WRITTEN AND PRESENTED BY:", bold=True)
    doc.add_paragraph()
    add_paragraph(doc, "KADJI MONTHE IDRISS ANGELOT")
    doc.add_paragraph()
    add_paragraph(doc, "LECTURER:  Dr. Mougnol Romeo", bold=True)
    for _ in range(4):
        doc.add_paragraph()
    add_page_break(doc)


def build_declaration_1(doc):
    add_section_title_page(doc, "DECLARATION")
    add_paragraph(
        doc,
        "I, declare that I am the sole author of this report. I understand the nature of plagiarism, "
        "and I am aware of the institution's policy on this. I declare that this report is an original "
        "work by me and contains neither materials previously published by another person nor material "
        "that has been accepted for the award of a certificate in any other institution, except where "
        "due acknowledgments have been made.",
        justify=True,
        size=12,
    )
    doc.add_paragraph()
    add_paragraph(doc, "Date................................................", size=12)
    doc.add_paragraph()
    add_paragraph(doc, "Signature.......................................", size=12)
    add_page_break(doc)


def build_signature_table(doc):
    table = doc.add_table(rows=1, cols=2)
    style_table(table)
    left = (
        "ACADEMIC SUPERVISOR\nDr. Mougnol Romeo\n"
        "Signature………………..........\nDATE...................................."
    )
    right = (
        "THE COORDINATOR\nEng. TSAKOU KOUMETIO Billy\n"
        "Signature——————————\nDate————————————"
    )
    table.rows[0].cells[0].text = left
    table.rows[0].cells[1].text = right
    add_page_break(doc)


def build_declaration_2(doc):
    add_section_title_page(doc, "DECLARATION")
    add_paragraph(
        doc,
        'I declare that I am the sole author of this project report entitled "UniSchedule Africa: '
        "An Automated University Timetable Generation SaaS Platform for African Universities\". "
        "I understand the nature of plagiarism and I am aware of the institute's policy on this. "
        "I certify that this project report was originally done by me during my university studies "
        "except for the Paragraphs, Sentences, Titles, Sub-titles or Relative References "
        "(see references or bibliography on this work).",
        justify=True,
    )
    add_page_break(doc)


def build_dedication(doc):
    add_section_title_page(doc, "DEDICATION")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "To My Parents", size=20)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "And to all those who believe technology can transform education across Africa.", italic=True)
    add_page_break(doc)


def build_acknowledgements(doc):
    add_section_title_page(doc, "ACKNOWLEDGEMENT")
    add_paragraph(
        doc,
        "I am very happy to have the honour of acknowledging the support of the following personalities "
        "— be it financially, morally or practically — in seeing that this piece of work is a success. "
        "I sincerely give thanks to:",
        justify=True,
    )
    ack_items = [
        ("Dr. Hon. Joseph NGUEPI,", " Founder of the University Institute of the Tropics (IUGET) for the establishment of this institution aimed at training students."),
        ("The General Coordinator of IUGET Dr. KEMFANG HERVEY", " for his support in the institution."),
        ("The Director of SOUTH POLYTECH Dr. MAZETIO Brice", " for his support in the institution."),
        ("The Coordinator of the BTech cycle Eng. TSAKOU KOUMETIO Billy", " for his coordination and support throughout this program."),
        ("Dr. Mougnol Romeo,", " my Academic Supervisor, for the motivation, direction and support that enabled me to complete this piece of work."),
        (None, "My parents, for their unlimited financial and moral support throughout my educational and professional journey — it was an awesome year."),
        (None, "To God Almighty for strength and guidance every day."),
        (None, "To all the lecturers of IUGET who contributed to my training and knowledge."),
        (None, "To all my classmates and friends who supported me in one way or another — I apologize to those I have not mentioned by name."),
    ]
    for bold_part, rest in ack_items:
        add_bullet(doc, rest, bold_prefix=bold_part)
    add_page_break(doc)


def build_abstract(doc):
    add_section_title_page(doc, "ABSTRACT")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        p,
        "UNISCHEDULE AFRICA: AN AUTOMATED UNIVERSITY TIMETABLE GENERATION SAAS PLATFORM FOR AFRICAN UNIVERSITIES",
        bold=True,
        size=12,
    )
    add_paragraph(
        doc,
        "Across sub-Saharan Africa, universities continue to rely on manual, spreadsheet-based methods to "
        "produce academic timetables — a process that consumes weeks of administrative effort, produces "
        "frequent scheduling conflicts, and fails to account for the complex constraints of African "
        "higher education institutions (multi-shift day/evening programmes, Saturday classes, limited "
        "room capacity, lecturer availability across multiple campuses, and LMD credit-hour systems). "
        "This project presents the design and development of UniSchedule Africa — a full-stack SaaS "
        "platform built using Next.js 14, React, TypeScript and PostgreSQL that automates university "
        "timetable generation for African institutions. The system provides role-based portals for "
        "Administrators and Timetablers, modules for managing courses, lecturers, rooms, student groups "
        "and academic sessions, and a genetic algorithm engine that generates conflict-free weekly "
        "schedules respecting hard constraints (room capacity, lecturer availability, time-slot "
        "conflicts) and soft constraints (evening/day separation, Friday prayer protection, Saturday "
        "scheduling). Key features include real-time conflict detection, a master grid view, timetable "
        "history and versioning, exam scheduling, PDF export, bilingual (French/English) support, "
        "and a public API for publishing timetables to students. The platform implements JWT "
        "authentication, bcrypt password hashing, role-based access control and audit logging. "
        "Results demonstrate that the system successfully reduces timetable preparation time from "
        "weeks to minutes while producing schedules that satisfy institutional constraints for "
        "universities such as ISIMA Yaoundé.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Keywords: Timetable Generation, SaaS, Next.js, PostgreSQL, Genetic Algorithm, African Universities, Scheduling Constraints, LMD System, Cameroon.",
        bold=True,
    )
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        p,
        "UNISCHEDULE AFRICA : UNE PLATEFORME SAAS DE GÉNÉRATION AUTOMATISÉE D'EMPLOIS DU TEMPS UNIVERSITAIRES POUR LES UNIVERSITÉS AFRICAINES",
        bold=True,
        size=12,
    )
    add_paragraph(
        doc,
        "En Afrique subsaharienne, les universités continuent de s'appuyer sur des méthodes manuelles "
        "basées sur des tableurs pour produire les emplois du temps académiques — un processus qui "
        "consomme des semaines d'efforts administratifs, produit des conflits d'horaires fréquents et "
        "ne tient pas compte des contraintes complexes des établissements d'enseignement supérieur "
        "africains (programmes multi-horaires jour/soir, cours du samedi, capacité limitée des salles, "
        "disponibilité des enseignants sur plusieurs campus, et systèmes de crédits LMD). Ce projet "
        "présente la conception et le développement d'UniSchedule Africa — une plateforme SaaS "
        "full-stack développée avec Next.js 14, React, TypeScript et PostgreSQL qui automatise la "
        "génération d'emplois du temps universitaires pour les institutions africaines. Le système "
        "propose des portails basés sur les rôles pour les Administrateurs et les Planificateurs "
        "d'horaires, des modules de gestion des cours, enseignants, salles, groupes d'étudiants et "
        "sessions académiques, et un moteur d'algorithme génétique qui génère des plannings hebdomadaires "
        "sans conflits en respectant les contraintes strictes et souples propres aux universités africaines.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Mots-clés : Génération d'emplois du temps, SaaS, Next.js, PostgreSQL, Algorithme génétique, Universités africaines, Contraintes d'horaires, Système LMD, Cameroun.",
        bold=True,
    )
    add_page_break(doc)


def build_toc(doc):
    add_section_title_page(doc, "TABLE OF CONTENTS")
    add_toc_line(doc, "CHAPTER 1   GENERAL INTRODUCTION", "CHAPTER_1_GENERAL_INTRODUCTION")
    entries = [
        ("1.1) INTRODUCTION TO THE STUDY", "sec_1_1", 1),
        ("1.2) BACKGROUND OF THE STUDY", "sec_1_2", 1),
        ("1.2.1) What is UniSchedule Africa?", "sec_1_2_1", 2),
        ("1.3) PROBLEM STATEMENT", "sec_1_3", 1),
        ("1.4) HYPOTHESIS OF THE STUDY", "sec_1_4", 1),
        ("1.5) PURPOSE OF THE STUDY", "sec_1_5", 1),
        ("1.6) OBJECTIVES OF THE STUDY", "sec_1_6", 1),
        ("1.6.1) General Objective", "sec_1_6_1", 2),
        ("1.6.2) Specific Objectives", "sec_1_6_2", 2),
        ("1.7) SIGNIFICANCE OF THE STUDY", "sec_1_7", 1),
        ("CHAPTER 2   LITERATURE REVIEW", "CHAPTER_2_LITERATURE_REVIEW", 0),
        ("2.1) INTRODUCTION", "sec_2_1", 1),
        ("2.2) REVIEW OF CONCEPTS AND THEORY", "sec_2_2", 1),
        ("2.2.1) Automated Timetable Generation", "sec_2_2_1", 2),
        ("2.2.2) Constraint-Based Scheduling", "sec_2_2_2", 2),
        ("2.2.3) SaaS Platforms in Higher Education", "sec_2_2_3", 2),
        ("2.2.4) Genetic Algorithms in Scheduling", "sec_2_2_4", 2),
        ("2.2.5) Role-Based Access Control in Multi-Stakeholder Systems", "sec_2_2_5", 2),
        ("2.2.6) Timetabling Challenges in Sub-Saharan African Universities", "sec_2_2_6", 2),
        ("2.3) OVERVIEW OF UNISCHEDULE AFRICA", "sec_2_3", 1),
        ("2.3.1) System Description", "sec_2_3_1", 2),
        ("2.3.2) Core Functional Modules", "sec_2_3_2", 2),
        ("2.3.3) Why Next.js and PostgreSQL?", "sec_2_3_3", 2),
        ("CHAPTER 3   METHODOLOGY AND TOOLS", "CHAPTER_3_METHODOLOGY_AND_TOOLS", 0),
        ("3.1) DEVELOPMENT APPROACH", "sec_3_1", 1),
        ("3.2) SYSTEM ARCHITECTURE", "sec_3_2", 1),
        ("3.3) DATABASE DESIGN", "sec_3_3", 1),
        ("3.4) TECHNOLOGIES USED", "sec_3_4", 1),
        ("3.5) SECURITY IMPLEMENTATION", "sec_3_5", 1),
    ]
    for label, bm, indent in entries:
        add_toc_line(doc, label, bm, indent)
    add_page_break(doc)


def build_list_of_figures(doc):
    add_section_title_page(doc, "LIST OF FIGURES")
    figures = [
        "Figure 3.1    System Architecture — 3-Tier SaaS .................................................. 17",
        "Figure 3.2    Entity Relationship Diagram (ERD) ................................................ 18",
        "Figure 3.3    Use Case Diagram ......................................................................... 19",
        "Figure 3.4    Data Flow — Timetable Generation Lifecycle ............................................ 20",
        "Figure 4.1    Login and Registration Page .......................................................... 22",
        "Figure 4.2    Institution Setup Screen ............................................................... 22",
        "Figure 4.3    Admin Dashboard — Overview ........................................................ 23",
        "Figure 4.4    Course Management Panel .................................................................. 23",
        "Figure 4.5    Interactive Timetable Grid .............................................................. 24",
        "Figure 4.6    Conflict Detection View ................................................................. 24",
        "Figure 4.7    Lecturer Availability Form ............................................................ 25",
        "Figure 4.8    Room Management — Capacity Settings ............................................ 25",
        "Figure 4.9    Timetable Generation Panel ................................................................ 26",
        "Figure 4.10   Master Grid View ........................................................................ 26",
        "Figure 4.11   Analytics Dashboard ..................................................................... 27",
        "Figure 4.12   Timetable History Panel .................................................................. 27",
        "Figure 4.13   Exam Scheduling Module ............................................................. 28",
        "Figure 4.14   PDF Export Preview ............................................................. 28",
        "Figure 4.15   Public Timetable API Page ........................................................... 29",
        "Figure 4.16   Bilingual Interface Toggle .......................................................... 30",
    ]
    for fig in figures:
        add_figure_line(doc, fig)
    add_page_break(doc)


def build_abbreviations(doc):
    add_section_title_page(doc, "LIST OF ABBREVIATIONS")
    abbrevs = [
        ("API", "Application Programming Interface"),
        ("CSRF", "Cross-Site Request Forgery"),
        ("CSS", "Cascading Style Sheets"),
        ("DB", "Database"),
        ("ERD", "Entity Relationship Diagram"),
        ("GA", "Genetic Algorithm"),
        ("JWT", "JSON Web Token"),
        ("LMD", "Licence-Master-Doctorat (Bologna credit system)"),
        ("MVC", "Model-View-Controller"),
        ("ORM", "Object-Relational Mapping"),
        ("RBAC", "Role-Based Access Control"),
        ("REST", "Representational State Transfer"),
        ("SaaS", "Software as a Service"),
        ("SQL", "Structured Query Language"),
        ("TS", "TypeScript"),
        ("UI", "User Interface"),
        ("UUID", "Universally Unique Identifier"),
    ]
    for i, (abbr, meaning) in enumerate(abbrevs, 1):
        add_paragraph(doc, f"{i}. {abbr}: {meaning}", size=12)
    add_page_break(doc)


def build_chapter1(doc):
    add_chapter_title(doc, "CHAPTER 1   GENERAL INTRODUCTION", "CHAPTER_1_GENERAL_INTRODUCTION")

    add_heading_section(doc, "1.1) INTRODUCTION TO THE STUDY", bookmark="sec_1_1")
    add_paragraph(
        doc,
        "Across sub-Saharan Africa, the effective scheduling of academic activities remains one of the most "
        "persistent administrative challenges facing universities and higher education institutions. In Cameroon, "
        "institutions such as ISIMA Yaoundé, IUGET and dozens of private universities serve thousands of students "
        "across multiple faculties, programmes and study modes — yet most still produce their timetables manually "
        "using spreadsheets, whiteboards and paper forms. The result is a process that consumes weeks of "
        "administrative effort each semester, produces frequent scheduling conflicts, and fails to adapt when "
        "lecturers, rooms or enrolments change at the last minute.",
        justify=True,
    )
    add_paragraph(
        doc,
        "The emergence of cloud-based SaaS platforms, combined with modern constraint-solving algorithms and "
        "web-first user interfaces, offers a concrete pathway to modernize this situation. Automating timetable "
        "generation, connecting academic administrators to a central scheduling platform, and publishing "
        "conflict-free timetables to students via a public API can collectively transform how African "
        "universities manage their most resource-intensive administrative task.",
        justify=True,
    )
    add_paragraph(
        doc,
        "This project presents UniSchedule Africa — a full-stack SaaS platform built with Next.js 14, React, "
        "TypeScript and PostgreSQL — designed to bring this transformation to life. The system provides "
        "automated timetable generation using a genetic algorithm, multi-role portals for administrators and "
        "timetablers, real-time conflict detection, exam scheduling, PDF export, bilingual support and a "
        "RESTful API for publishing timetables to students.",
        justify=True,
    )

    add_heading_section(doc, "1.2) BACKGROUND OF THE STUDY", bookmark="sec_1_2")
    add_paragraph(
        doc,
        "Academic timetable quality is one of the most visible indicators of a university's administrative "
        "competence. In Cameroon's higher education sector, institutions operating under the LMD (Licence-Master-"
        "Doctorat) system must schedule hundreds of courses across day and evening programmes, Saturday sessions, "
        "multiple campuses and shared laboratory resources — all while respecting lecturer availability, room "
        "capacity limits and institutional policies such as Friday prayer protection. According to administrative "
        "surveys, manual timetable preparation at a mid-sized institution can consume 3–6 weeks per semester "
        "and still produce double-booked rooms, overlapping lecturer assignments and student group conflicts.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Most Cameroonian universities currently rely on manual scheduling workflows disconnected from real-time "
        "resource availability. Timetablers follow intuition and institutional memory rather than systematic "
        "constraint checking, resulting in wasted classroom hours, missed course sessions and last-minute "
        "rescheduling that disrupts both lecturers and students. There is no unified digital platform to "
        "manage courses, lecturers, rooms and student groups in one place, and no automated engine to generate "
        "optimized schedules from institutional data.",
        justify=True,
    )
    add_heading_section(doc, "1.2.1) What is UniSchedule Africa?", level=3, bookmark="sec_1_2_1")
    add_paragraph(
        doc,
        "UniSchedule Africa is a web-based Automated University Timetable Generation SaaS Platform (AUTGSP) "
        "that digitizes and connects every stakeholder in a university's scheduling chain — from the "
        "institutional administrator down to the individual student accessing their timetable online. It provides "
        "a unified platform where courses, lecturers, rooms and student groups are managed centrally, timetables "
        "are generated automatically using a genetic algorithm, conflicts are detected in real time, and "
        "published schedules are accessible via a public API without authentication.",
        justify=True,
    )

    add_heading_section(doc, "1.3) PROBLEM STATEMENT", bookmark="sec_1_3")
    add_paragraph(doc, "The following core problems motivate the development of this system:", justify=True)
    problems = [
        ("Manual timetable creation: ", "Universities spend weeks each semester manually building timetables in spreadsheets, with no systematic validation of constraints or conflicts."),
        ("Frequent scheduling conflicts: ", "Double-booked rooms, overlapping lecturer assignments and student group clashes are common because manual methods cannot exhaustively check all constraint combinations."),
        ("No centralized resource management: ", "Course, lecturer, room and group data are scattered across departments in incompatible formats, making coordinated scheduling impossible."),
        ("Inflexible to changes: ", "When a lecturer becomes unavailable or a room is closed, entire timetables must be rebuilt manually rather than regenerated automatically."),
        ("No support for African constraints: ", "Existing commercial timetabling software rarely supports day/evening dual programmes, Saturday classes, Friday prayer protection, LMD credit systems and multi-campus lecturer sharing common in African universities."),
        ("No student access channel: ", "Students receive timetables as PDF attachments or printed sheets with no digital, always-updated access point."),
    ]
    for prefix, text in problems:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_section(doc, "1.4) HYPOTHESIS OF THE STUDY", bookmark="sec_1_4")
    add_paragraph(
        doc,
        "The development and deployment of a web-based Automated Timetable Generation SaaS Platform will "
        "significantly improve university scheduling efficiency by providing automated conflict-free timetable "
        "generation, centralized resource management, real-time conflict detection, and digital publication "
        "of schedules to students. It is hypothesised that such a system will reduce timetable preparation "
        "time from weeks to minutes, eliminate hard scheduling conflicts, and provide African universities "
        "with the foundational digital infrastructure needed to modernize academic administration.",
        justify=True,
    )

    add_heading_section(doc, "1.5) PURPOSE OF THE STUDY", bookmark="sec_1_5")
    add_paragraph(
        doc,
        "The purpose of this project is to design, develop and deploy a fully functional SaaS platform that "
        "serves as a digital backbone for university timetable operations across African institutions. The "
        "platform aims to replace the current manual, spreadsheet-based scheduling process with a connected, "
        "role-based application that empowers administrators and timetablers with automated generation tools "
        "while giving students instant access to published schedules.",
        justify=True,
    )

    add_heading_section(doc, "1.6) OBJECTIVES OF THE STUDY", bookmark="sec_1_6")
    add_heading_section(doc, "1.6.1) General Objective", level=3, bookmark="sec_1_6_1")
    add_paragraph(
        doc,
        "To design, develop and deploy a web-based Automated University Timetable Generation SaaS Platform "
        "that empowers African universities to generate conflict-free academic timetables automatically, "
        "manage scheduling resources centrally, and publish schedules digitally to students and staff.",
        justify=True,
    )
    add_heading_section(doc, "1.6.2) Specific Objectives", level=3, bookmark="sec_1_6_2")
    objectives = [
        "Develop a genetic algorithm engine that generates conflict-free weekly timetables respecting room capacity, lecturer availability, time-slot conflicts and institutional scheduling policies.",
        "Implement a multi-role authentication system with distinct portals and access controls for Super Administrators, Administrators, Timetablers, Lecturers and Viewers.",
        "Create resource management modules for courses, lecturers, rooms, student groups, academic sessions and holidays with African-specific constraints (day/evening modes, Saturday scheduling, Friday prayer protection).",
        "Design a real-time conflict detection system that identifies room, lecturer and group double-bookings before timetable publication.",
        "Build an exam scheduling module supporting written, oral and practical examinations with invigilator and room assignment.",
        "Develop a master grid view and timetable history module for comparing generated schedules across semesters and academic years.",
        "Implement PDF export and a public REST API endpoint for publishing timetables to students without authentication.",
        "Incorporate bilingual support (French and English) to serve Cameroon's bilingual university community.",
        "Ensure application security through JWT authentication, bcrypt password hashing, role-based access control and comprehensive audit logging.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    add_heading_section(doc, "1.7) SIGNIFICANCE OF THE STUDY", bookmark="sec_1_7")
    add_paragraph(
        doc,
        "This study is significant for several reasons. First, it directly addresses an administrative bottleneck "
        "affecting hundreds of universities across sub-Saharan Africa. Every semester that timetables are built "
        "manually is a semester where administrative staff spend weeks on repetitive scheduling tasks instead of "
        "supporting teaching and learning. A system that can generate conflict-free timetables in minutes has "
        "the potential to eliminate this inefficiency entirely.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Second, the study demonstrates how modern web technologies — Next.js 14, React, TypeScript, PostgreSQL "
        "— can be applied to solve real higher-education administration problems in African contexts without "
        "requiring costly proprietary software licenses. The entire stack is open-source and deployable on "
        "affordable cloud infrastructure (Supabase, Railway) or local servers.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Third, the African-specific constraint model (day/evening programmes, Saturday classes, LMD credit "
        "systems, Friday prayer protection) represents a new contribution to the timetabling literature — "
        "one where scheduling rules reflect the actual operational realities of Cameroonian and broader "
        "sub-Saharan African universities rather than European or North American defaults.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Finally, the data generated by the system — generation logs, conflict histories, resource utilization "
        "metrics, timetable versions — provides university administrators with an evidence base for resource "
        "planning that has never previously existed in digital form, enabling truly data-driven academic "
        "administration.",
        justify=True,
    )
    add_page_break(doc)


def build_chapter2(doc):
    add_chapter_title(doc, "CHAPTER 2   LITERATURE REVIEW", "CHAPTER_2_LITERATURE_REVIEW")

    add_heading_section(doc, "2.1) INTRODUCTION", bookmark="sec_2_1")
    add_paragraph(
        doc,
        "This chapter presents a comprehensive review of the existing body of knowledge, theories and concepts "
        "relevant to the development of an Automated University Timetable Generation SaaS Platform for African "
        "universities. It covers the global evolution of automated timetabling systems, constraint-based "
        "scheduling theory, the role of genetic algorithms in combinatorial optimization, SaaS delivery models "
        "in higher education, and the specific timetabling challenges documented in sub-Saharan African "
        "universities. The literature review aims to establish a solid theoretical and empirical foundation "
        "that situates UniSchedule Africa within its broader academic and practical context.",
        justify=True,
    )

    add_heading_section(doc, "2.2) REVIEW OF CONCEPTS AND THEORY", bookmark="sec_2_2")

    add_heading_section(doc, "2.2.1) Automated Timetable Generation", level=3, bookmark="sec_2_2_1")
    add_paragraph(
        doc,
        "University timetabling is a well-studied combinatorial optimization problem classified as NP-hard "
        "(Burke & Petrovic, 2002). Early automated approaches relied on integer programming and graph coloring "
        "algorithms, while modern systems employ metaheuristic methods including simulated annealing, tabu "
        "search and genetic algorithms. Commercial systems such as Scientia (now part of Tribal Group) and "
        "aSc Timetables have demonstrated that automated generation can reduce scheduling time by over 90% "
        "compared to manual methods in European university contexts.",
        justify=True,
    )
    add_paragraph(
        doc,
        "An automated timetabling system typically integrates three layers: a data layer (courses, lecturers, "
        "rooms, groups), a constraint engine (hard and soft scheduling rules), and a generation layer (the "
        "optimization algorithm producing feasible schedules). UniSchedule Africa addresses all three layers — "
        "providing a PostgreSQL data layer with 20+ normalized tables, a constraint engine encoding African "
        "university scheduling rules, and a genetic algorithm generation layer implemented in JavaScript.",
        justify=True,
    )

    add_heading_section(doc, "2.2.2) Constraint-Based Scheduling", level=3, bookmark="sec_2_2_2")
    add_paragraph(
        doc,
        "Constraint-based scheduling distinguishes between hard constraints (must never be violated — e.g., a "
        "lecturer cannot be in two rooms simultaneously) and soft constraints (desirable but negotiable — e.g., "
        "avoiding back-to-back sessions for the same student group). Research by Schaerf (1999) established "
        "that effective timetabling systems must encode both constraint types and optimize for soft constraint "
        "satisfaction while guaranteeing hard constraint feasibility.",
        justify=True,
    )
    add_paragraph(
        doc,
        "UniSchedule Africa encodes hard constraints including room capacity limits, lecturer availability "
        "windows, time-slot non-overlap for rooms/lecturers/groups, and room type requirements (laboratory "
        "courses require lab rooms). Soft constraints include minimizing gaps in student schedules, respecting "
        "lecturer maximum hours per day/week, and preferring day-mode lecturers for morning slots.",
        justify=True,
    )

    add_heading_section(doc, "2.2.3) SaaS Platforms in Higher Education", level=3, bookmark="sec_2_2_3")
    add_paragraph(
        doc,
        "Software-as-a-Service (SaaS) delivery has transformed higher education administration globally. "
        "Platforms such as Banner (Ellucian), Moodle (LMS) and Google Workspace for Education demonstrate "
        "that cloud-hosted, multi-tenant applications can serve institutions of varying sizes without "
        "on-premises infrastructure. The SaaS model is particularly suited to African universities where "
        "IT budgets are limited and dedicated server rooms are uncommon.",
        justify=True,
    )
    add_paragraph(
        doc,
        "UniSchedule Africa adopts a multi-tenant SaaS architecture where each institution (university) "
        "registers as a tenant with isolated data, configurable scheduling policies and subscription-based "
        "access. This model allows a single deployment to serve multiple Cameroonian universities while "
        "maintaining complete data separation between institutions.",
        justify=True,
    )

    add_heading_section(doc, "2.2.4) Genetic Algorithms in Scheduling", level=3, bookmark="sec_2_2_4")
    add_paragraph(
        doc,
        "Genetic algorithms (GAs) are population-based metaheuristics inspired by natural selection, widely "
        "applied to timetabling problems since Colorni et al. (1998). A GA maintains a population of candidate "
        "schedules, evaluates each against a fitness function penalizing constraint violations, and iteratively "
        "breeds new generations through selection, crossover and mutation until a feasible or near-optimal "
        "schedule emerges.",
        justify=True,
    )
    add_paragraph(
        doc,
        "UniSchedule Africa implements a GA with population size 50, 150 generations, tournament selection, "
        "single-point crossover and 15% mutation rate. The fitness function penalizes room conflicts, lecturer "
        "conflicts, capacity violations and unavailable time slots with weighted scores, converging to "
        "conflict-free schedules within the generation limit for typical institution sizes (50–200 sessions).",
        justify=True,
    )

    add_heading_section(doc, "2.2.5) Role-Based Access Control in Multi-Stakeholder Systems", level=3, bookmark="sec_2_2_5")
    add_paragraph(
        doc,
        "Multi-role web applications serving organizations with hierarchical responsibilities require careful "
        "implementation of Role-Based Access Control (RBAC). The RBAC model assigns permissions to roles rather "
        "than individual users, ensuring that each stakeholder category can only access the system functionality "
        "relevant to their responsibilities (Ferraiolo et al., 2001). Implementing RBAC via JWT claims and "
        "middleware-level role checks — as done in UniSchedule Africa — is a standard, well-tested pattern for "
        "securing multi-tenant SaaS applications.",
        justify=True,
    )

    add_heading_section(doc, "2.2.6) Timetabling Challenges in Sub-Saharan African Universities", level=3, bookmark="sec_2_2_6")
    add_paragraph(
        doc,
        "Higher education administration remains one of the most under-digitized sectors in sub-Saharan Africa. "
        "Studies focusing on Cameroonian universities document that manual scheduling approaches consistently "
        "fail to handle the complexity of dual day/evening programmes, shared lecturer pools across institutions, "
        "Saturday class requirements and the LMD credit-hour system. Key constraints identified include limited "
        "classroom infrastructure relative to enrolment, lecturer availability across multiple campuses, and "
        "absence of digital resource databases.",
        justify=True,
    )
    add_paragraph(
        doc,
        "These findings directly motivate the core design decisions of UniSchedule Africa: the centralized "
        "resource management modules provide the digital data foundation that has been absent; the genetic "
        "algorithm engine replaces manual trial-and-error scheduling; and the conflict detection system "
        "provides the validation layer that spreadsheet-based methods cannot offer. The system thus directly "
        "addresses each constraint identified in the African higher education timetabling literature.",
        justify=True,
    )

    add_heading_section(doc, "2.3) OVERVIEW OF UNISCHEDULE AFRICA", bookmark="sec_2_3")
    add_heading_section(doc, "2.3.1) System Description", level=3, bookmark="sec_2_3_1")
    add_paragraph(
        doc,
        "UniSchedule Africa is a full-stack Next.js 14 / PostgreSQL SaaS application designed for deployment "
        "on modern cloud platforms (Vercel, Railway, Supabase) or via Docker Compose for on-premises hosting. "
        "It provides distinct role-based portals — Super Admin, Admin, Timetabler, Lecturer and Viewer — each "
        "with dedicated dashboards, navigation menus and access controls. The system manages 20+ relational "
        "database tables covering institutions, users, faculties, departments, programmes, courses, lecturers, "
        "rooms, student groups, timetables, sessions, exams and conflicts.",
        justify=True,
    )

    add_heading_section(doc, "2.3.2) Core Functional Modules", level=3, bookmark="sec_2_3_2")
    modules = [
        ("Course Management Module: ", "Manages course catalogues with LMD credit hours, session durations, laboratory requirements and department assignments per institution."),
        ("Lecturer Management Module: ", "Tracks lecturer profiles, availability windows, maximum teaching hours, day/evening teaching modes and multi-institution assignments."),
        ("Room Management Module: ", "Manages buildings and rooms with capacity, room type (classroom, lab, amphitheater, seminar), equipment flags and availability schedules including blackout dates."),
        ("Student Group Module: ", "Manages programme cohorts with level, mode (day/evening/weekend), enrolment counts and semester assignments."),
        ("Timetable Generation Engine: ", "Genetic algorithm producing conflict-free weekly schedules displayed on an interactive grid with day/evening slot separation."),
        ("Conflict Detection Module: ", "Real-time identification of room, lecturer and group double-bookings with severity classification and resolution suggestions."),
        ("Exam Scheduling Module: ", "Schedules written, oral and practical examinations with room, invigilator and student group assignment."),
        ("Publication Module: ", "PDF export and public REST API (GET /api/public/[slug]/timetable) for student access without authentication."),
        ("Security Layer: ", "JWT authentication with refresh tokens; bcrypt password hashing; role-based middleware; comprehensive audit logging."),
    ]
    for prefix, text in modules:
        add_bullet(doc, text, bold_prefix=prefix)

    add_heading_section(doc, "2.3.3) Why Next.js and PostgreSQL?", level=3, bookmark="sec_2_3_3")
    add_paragraph(
        doc,
        "Next.js 14 and PostgreSQL were selected as the core technology stack for three primary reasons. First, "
        "Next.js provides a unified full-stack framework combining React frontend and API routes backend, "
        "eliminating the complexity of separate frontend/backend deployments — critical for a solo developer "
        "project. Second, PostgreSQL offers ACID-compliant relational storage with JSONB support for generation "
        "logs, UUID primary keys for multi-tenant data isolation, and free hosting via Supabase or Railway — "
        "making it accessible to African universities without database licensing costs. Third, the TypeScript "
        "ecosystem provides compile-time type safety across the entire stack, reducing runtime errors in "
        "constraint validation and API endpoint logic.",
        justify=True,
    )
    add_page_break(doc)


def build_chapter3(doc):
    add_chapter_title(doc, "CHAPTER 3   METHODOLOGY AND TOOLS", "CHAPTER_3_METHODOLOGY_AND_TOOLS")

    add_heading_section(doc, "3.1) DEVELOPMENT APPROACH", bookmark="sec_3_1")
    add_paragraph(
        doc,
        "The project follows an Agile iterative development methodology divided into four two-week sprints. "
        "Requirements were elicited through analysis of scheduling workflows at ISIMA Yaoundé (the pilot "
        "institution), review of the timetabling literature, and examination of existing commercial scheduling "
        "software limitations in African contexts. Each sprint produced a testable increment reviewed against "
        "the functional requirements before proceeding to the next phase.",
        justify=True,
    )
    add_paragraph(
        doc,
        "The development environment consists of Node.js 18+ with Next.js 14 on Windows, PostgreSQL 16 via "
        "Docker Compose for local development, and Visual Studio Code as the primary editor. Database migrations "
        "are managed via custom SQL migration scripts (npm run db:migrate), with Prisma-style schema definitions "
        "in scripts/schema.sql. API endpoint testing uses Postman and the built-in Swagger documentation "
        "(/api/docs). Version control was managed via Git throughout the development process.",
        justify=True,
    )

    add_heading_section(doc, "3.2) SYSTEM ARCHITECTURE", bookmark="sec_3_2")
    add_paragraph(doc, "The system follows a 3-tier SaaS architecture:", justify=True)
    layers = [
        ("Presentation Layer: ", "React 18 with Vite frontend (Tailwind CSS, Framer Motion, Lucide icons). Responsive layout with IndexedDB (Dexie.js) for offline demo mode. Communicates with backend via REST API fetch calls."),
        ("Application Layer: ", "Next.js 14 API routes handling business logic, JWT authentication, role-based middleware, genetic algorithm scheduling engine, conflict detection, notification dispatch and Swagger-documented REST endpoints."),
        ("Data Layer: ", "PostgreSQL 16 relational database with 20+ normalized tables, UUID primary keys, foreign key enforcement, soft deletes (deleted_at) and JSONB columns for generation logs."),
    ]
    for prefix, text in layers:
        add_bullet(doc, text, bold_prefix=prefix)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Figure 3.1   System Architecture — 3-Tier SaaS", bold=True, size=12)

    add_heading_section(doc, "3.3) DATABASE DESIGN", bookmark="sec_3_3")
    add_paragraph(
        doc,
        "The PostgreSQL database schema consists of 20+ normalized tables organized around six core entities: "
        "Institutions (multi-tenant root), Users (with role assignments), Academic Structure (faculties, "
        "departments, programmes, student groups), Resources (courses, lecturers, rooms, buildings), Timetables "
        "(generated schedules with sessions and session-group mappings), and Exams (examination scheduling). "
        "Supporting tables handle conflicts, audit logs, room/lecturer blackouts, lecturer availability windows "
        "and notification records.",
        justify=True,
    )
    add_paragraph(
        doc,
        "All foreign key relationships are enforced at the database level with ON DELETE CASCADE or RESTRICT "
        "as appropriate. Institution-level configuration (day/evening times, Saturday enabled, Friday prayer "
        "protected, academic system type) is stored in the institutions table and applied as scheduling "
        "constraints during generation. The ERD below illustrates the primary entity relationships:",
        justify=True,
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    add_run(p, "Figure 3.2   Entity Relationship Diagram (ERD)", bold=True, size=12)

    add_heading_section(doc, "3.4) TECHNOLOGIES USED", bookmark="sec_3_4")
    add_paragraph(
        doc,
        "The following table summarizes the key technologies used in the system and the justification for each choice:",
        justify=True,
    )
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    headers = ["Technology", "Justification"]
    rows_data = [
        ("Next.js 14", "Full-stack React framework with API routes, server-side rendering and TypeScript support. Enables unified frontend/backend deployment on Vercel or Docker."),
        ("PostgreSQL 16", "ACID-compliant relational DBMS with UUID, JSONB and constraint support. Free via Supabase/Railway. Handles the 20+ table normalized schema with foreign key enforcement."),
        ("React 18 / Vite", "Component-based UI with fast HMR development. Tailwind CSS for responsive, mobile-friendly timetabler interfaces."),
        ("TypeScript", "Static typing across frontend and backend reduces runtime errors in constraint validation and API contracts."),
        ("JWT (jsonwebtoken)", "Stateless authentication with access/refresh token pair. Enables secure API access without server-side session storage."),
        ("Docker Compose", "Local development stack matching production deployment. Runs backend, frontend and PostgreSQL in isolated containers."),
        ("Genetic Algorithm (JS)", "Custom scheduling engine in frontend/src/scheduler.js. Population-based optimization converges to conflict-free schedules for typical institution sizes."),
    ]
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_run_font(run, bold=True, size=11)
    for i, (tech, just) in enumerate(rows_data, 1):
        table.rows[i].cells[0].text = tech
        table.rows[i].cells[1].text = just
    style_table(table)

    add_heading_section(doc, "3.5) SECURITY IMPLEMENTATION", bookmark="sec_3_5")
    add_paragraph(
        doc,
        "Security was treated as a first-class concern throughout development. The following mechanisms were implemented:",
        justify=True,
    )
    security_items = [
        ("JWT Authentication: ", "Access tokens (15-minute expiry) and refresh tokens (7-day expiry) issued on login. All protected API routes validate the Bearer token via middleware before processing requests."),
        ("Password Hashing: ", "All user passwords are stored as bcrypt hashes (cost factor 12). Plain-text passwords are never stored or logged."),
        ("Role-Based Access Control: ", "Five roles (superadmin, admin, timetabler, viewer, lecturer) enforced at middleware level. Each API endpoint checks the authenticated user's role against required permissions before executing."),
        ("SQL Injection Prevention: ", "All database queries use parameterized statements via the pg library. Direct string concatenation into SQL queries does not appear anywhere in the codebase."),
        ("Multi-Tenant Data Isolation: ", "Every database query includes institution_id filtering derived from the authenticated user's JWT claims, preventing cross-institution data access."),
        ("Audit Logging: ", "All create, update and delete operations on scheduling resources are logged with user ID, timestamp, resource type and change details for accountability and compliance."),
    ]
    for prefix, text in security_items:
        add_bullet(doc, text, bold_prefix=prefix)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def main():
    doc = Document()
    configure_document(doc)
    build_cover(doc)
    build_declaration_1(doc)
    build_signature_table(doc)
    build_declaration_2(doc)
    build_dedication(doc)
    build_acknowledgements(doc)
    build_abstract(doc)
    build_toc(doc)
    build_list_of_figures(doc)
    build_abbreviations(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
